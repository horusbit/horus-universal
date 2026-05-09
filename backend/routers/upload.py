"""
Router de Upload — HORUS Universal
Procesa PDF, DOCX, XLSX, TXT e imágenes para inyectarlos como contexto al agente.
"""
import os
import io
import uuid
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from fastapi.responses import JSONResponse
from auth.supabase_auth import get_optional_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/upload", tags=["upload"])

MAX_SIZE_MB = 10
ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/msword": "doc",
    "text/plain": "txt",
    "text/csv": "csv",
    "image/jpeg": "image",
    "image/png": "image",
    "image/webp": "image",
    "image/gif": "image",
}


@router.post("/")
async def upload_file(
    file: UploadFile = File(...),
    user=Depends(get_optional_user),
):
    """Procesa un archivo y retorna su contenido como texto para el agente."""
    # Validar tipo
    content_type = file.content_type or ""
    file_type = ALLOWED_TYPES.get(content_type)
    if not file_type:
        # Try by extension
        name = file.filename or ""
        if name.endswith(".pdf"):      file_type = "pdf"
        elif name.endswith(".docx"):   file_type = "docx"
        elif name.endswith(".xlsx"):   file_type = "xlsx"
        elif name.endswith(".txt"):    file_type = "txt"
        elif name.endswith(".csv"):    file_type = "csv"
        elif name.endswith((".jpg", ".jpeg", ".png", ".webp")): file_type = "image"
        else:
            raise HTTPException(400, f"Tipo de archivo no soportado: {content_type}")

    # Validar tamaño
    data = await file.read()
    size_mb = len(data) / (1024 * 1024)
    if size_mb > MAX_SIZE_MB:
        raise HTTPException(400, f"Archivo muy grande ({size_mb:.1f}MB). Máximo {MAX_SIZE_MB}MB.")

    filename = file.filename or f"archivo_{uuid.uuid4().hex[:8]}"
    extracted = ""
    error_msg = ""

    try:
        if file_type == "pdf":
            extracted = _extract_pdf(data)
        elif file_type == "docx":
            extracted = _extract_docx(data)
        elif file_type == "xlsx":
            extracted = _extract_xlsx(data)
        elif file_type in ("txt", "csv"):
            extracted = data.decode("utf-8", errors="replace")
        elif file_type == "image":
            # Para imágenes retornamos base64 — el LLM lo puede ver
            import base64
            b64 = base64.b64encode(data).decode()
            return JSONResponse({
                "filename": filename,
                "type": "image",
                "content_type": content_type,
                "base64": b64,
                "size_mb": round(size_mb, 2),
                "context": f"[ARCHIVO IMAGEN: {filename}]",
            })
    except Exception as e:
        logger.error(f"[Upload] Error procesando {filename}: {e}")
        error_msg = str(e)

    # Truncar si es muy largo (max ~8000 chars = ~2000 tokens)
    MAX_CHARS = 8000
    truncated = False
    if len(extracted) > MAX_CHARS:
        extracted = extracted[:MAX_CHARS]
        truncated = True

    if not extracted and not error_msg:
        error_msg = "No se pudo extraer texto del archivo."

    context = build_file_context(filename, file_type, extracted, truncated)

    return JSONResponse({
        "filename": filename,
        "type": file_type,
        "size_mb": round(size_mb, 2),
        "chars_extracted": len(extracted),
        "truncated": truncated,
        "context": context,
        "error": error_msg or None,
    })


def _extract_pdf(data: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        texts = []
        for i, page in enumerate(reader.pages[:30]):  # max 30 páginas
            t = page.extract_text() or ""
            if t.strip():
                texts.append(f"--- Página {i+1} ---\n{t.strip()}")
        return "\n\n".join(texts)
    except Exception as e:
        # Fallback: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n\n".join(
                    p.extract_text() or "" for p in pdf.pages[:30]
                ).strip()
        except Exception:
            raise e


def _extract_docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # También tablas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _extract_xlsx(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for sheet in wb.worksheets[:5]:  # max 5 hojas
            lines.append(f"=== Hoja: {sheet.title} ===")
            for row in sheet.iter_rows(max_row=200, values_only=True):
                row_text = " | ".join(str(v) for v in row if v is not None)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception:
        # Fallback pandas
        import pandas as pd
        dfs = pd.read_excel(io.BytesIO(data), sheet_name=None, nrows=200)
        parts = []
        for name, df in list(dfs.items())[:5]:
            parts.append(f"=== Hoja: {name} ===\n{df.to_string(index=False, max_rows=100)}")
        return "\n\n".join(parts)


def build_file_context(filename: str, file_type: str, content: str, truncated: bool) -> str:
    type_labels = {
        "pdf": "📄 PDF", "docx": "📝 Word", "xlsx": "📊 Excel",
        "txt": "📃 Texto", "csv": "📊 CSV", "image": "🖼️ Imagen",
    }
    label = type_labels.get(file_type, "📎 Archivo")
    trunc_note = " (truncado — archivo muy largo)" if truncated else ""
    header = f"[{label} ADJUNTO: {filename}{trunc_note}]\nContenido del archivo:\n\n"
    footer = "\n\n[FIN DEL ARCHIVO — Responde basándote en este contenido]"
    return header + (content or "(sin contenido extraíble)") + footer
